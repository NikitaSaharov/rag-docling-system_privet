#!/usr/bin/env python3
"""
Скрипт для создания Word документа со стандартом тестирования RAG-системы
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Ошибка: установите библиотеку python-docx")
    print("Выполните: pip install python-docx")
    exit(1)

def add_heading(doc, text, level=1):
    """Добавляет заголовок"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, size=11):
    """Добавляет параграф"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    return p

def add_code_block(doc, text):
    """Добавляет блок кода"""
    p = doc.add_paragraph(text)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 128)
    p.paragraph_format.left_indent = Inches(0.5)
    return p

def create_document():
    """Создает Word документ со стандартом"""
    doc = Document()
    
    # Заголовок документа
    title = doc.add_heading('Стандарт "Вопрос-Ответ" для тестирования RAG-системы', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 1. Параметры векторизации
    add_heading(doc, '1. Параметры векторизации системы', 1)
    
    params = [
        ('Модель эмбеддингов:', 'nomic-embed-text (через Ollama)'),
        ('Размер чанка:', '350 слов'),
        ('Перекрытие (overlap):', '70 слов'),
        ('Размерность вектора:', 'определяется моделью nomic-embed-text')
    ]
    
    for key, value in params:
        p = doc.add_paragraph()
        p.add_run(key).bold = True
        p.add_run(f' {value}')
    
    doc.add_paragraph()
    
    # 2. Параметры генерации ответов
    add_heading(doc, '2. Параметры генерации ответов', 1)
    
    gen_params = [
        ('Модель:', 'DeepSeek Chat (через Polza.ai API)'),
        ('Temperature:', '0.0 (максимальная точность для формул)'),
        ('Top-p:', '0.95'),
        ('Max tokens:', '2000'),
        ('Таймаут:', '60 секунд')
    ]
    
    for key, value in gen_params:
        p = doc.add_paragraph()
        p.add_run(key).bold = True
        p.add_run(f' {value}')
    
    doc.add_paragraph()
    
    # 3. Алгоритм поиска
    add_heading(doc, '3. Алгоритм поиска (Hybrid Search)', 1)
    
    search_params = [
        ('Semantic search:', 'векторный поиск по косинусной близости'),
        ('Keyword boosting:', '+0.3 за совпадение документа (Справочник, Золотой Стандарт)'),
        ('Formula boosting:', '+0.25 за наличие формул в чанке'),
        ('Definition boosting:', '+0.5 за чанки с определениями'),
        ('Small document boost:', '+0.05 для документов <100 чанков'),
        ('Score threshold:', '0.60 (минимальный порог релевантности)'),
        ('Context expansion:', '±1 соседний чанк для формул')
    ]
    
    for key, value in search_params:
        p = doc.add_paragraph()
        p.add_run(key).bold = True
        p.add_run(f' {value}')
    
    doc.add_paragraph()
    
    # 4. Стандартный формат тестирования
    add_heading(doc, '4. Стандартный формат тестирования', 1)
    add_heading(doc, 'Типы вопросов для проверки:', 2)
    
    question_types = [
        ('A. Определения (дефиниции)', ['Что такое [термин]?', 'Определение [термина]', 'Что это [термин]?']),
        ('B. Формулы и расчёты', ['Как рассчитать [показатель]?', 'Формула для [показателя]', 'Как считать [метрику]?']),
        ('C. Практические задачи', ['Что делать если [проблема]?', 'Как улучшить [показатель]?', 'Как решить [ситуацию]?']),
        ('D. Расчёты с данными', ['Посчитай [показатель] при данных: [числа]'])
    ]
    
    for type_name, examples in question_types:
        p = doc.add_paragraph()
        p.add_run(type_name).bold = True
        for ex in examples:
            doc.add_paragraph(f'- {ex}', style='List Bullet')
    
    doc.add_page_break()
    
    # 5. ПРИМЕРЫ СТАНДАРТНЫХ ТЕСТОВ
    add_heading(doc, '5. ПРИМЕРЫ СТАНДАРТНЫХ ТЕСТОВ', 1)
    
    # Тест 1
    add_heading(doc, 'Тест №1: Определение (базовый)', 2)
    
    p = doc.add_paragraph()
    p.add_run('Вопрос:').bold = True
    add_code_block(doc, 'Что такое нормочас?')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Ожидаемый ответ (структура):').bold = True
    
    answer_structure = """[Определение из документа дословно]

Где:
- [переменная 1] - расшифровка
- [переменная 2] - расшифровка

Нормы:
- [категория 1]: [значения]
- [категория 2]: [значения]

Вопросы:
1. [уточняющий вопрос 1]
2. [уточняющий вопрос 2]"""
    
    add_code_block(doc, answer_structure)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Критерии проверки:').bold = True
    
    criteria = [
        '✅ Дословная цитата определения из документа',
        '✅ Формула указана без markdown (**)',
        '✅ Все переменные расшифрованы',
        '✅ Указаны нормативные значения',
        '✅ 2-3 уточняющих вопроса',
        '❌ НЕТ выдуманной информации',
        '❌ НЕТ фраз "В документе", "Согласно контексту"'
    ]
    
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    
    doc.add_paragraph()
    
    # Тест 2
    add_heading(doc, 'Тест №2: Формула (расширенный)', 2)
    
    p = doc.add_paragraph()
    p.add_run('Вопрос:').bold = True
    add_code_block(doc, 'Как рассчитать валовую выручку клиники?')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Ожидаемый ответ (структура):').bold = True
    
    answer_structure2 = """ВВклиники = Кзаг * НЧ * tраб * n

Где:
- Кзаг - коэффициент загрузки врачей
- НЧ - средний нормочас врачей
- tраб - рабочее время врачей
- n - количество врачей

Для расчёта нужны данные:
1. [список данных]

Пример расчёта:
[пошаговый расчёт, если есть данные]

Вопросы:
1. Как рассчитать коэффициент загрузки?
2. Какие нормы нормочаса для терапевтов?"""
    
    add_code_block(doc, answer_structure2)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Критерии проверки:').bold = True
    
    criteria2 = [
        '✅ Формула указана точно (НЕ в markdown)',
        '✅ Каждая переменная объяснена',
        '✅ Указаны требуемые данные для расчёта',
        '✅ Приведены нормы (если есть в документе)',
        '❌ НЕТ подстановки выдуманных чисел'
    ]
    
    for criterion in criteria2:
        doc.add_paragraph(criterion, style='List Bullet')
    
    doc.add_page_break()
    
    # Тест 3
    add_heading(doc, 'Тест №3: Расчёт с данными (проверка вычислений)', 2)
    
    p = doc.add_paragraph()
    p.add_run('Вопрос:').bold = True
    add_code_block(doc, 'Посчитай нормочас врача, если его валовая выручка 600 000 рублей, а время работы с пациентами 80 часов.')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Ожидаемый ответ (структура):').bold = True
    
    answer_structure3 = """Формула: НЧ = ВВ / tзаг

Подстановка данных:
НЧ = 600 000 руб / 80 часов

Расчёт:
НЧ = 7 500 руб/час

Интерпретация:
[сравнение с нормами для разных специализаций]

Вопросы:
1. Как увеличить нормочас врача?
2. Влияет ли нормочас на зарплату?"""
    
    add_code_block(doc, answer_structure3)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Критерии проверки:').bold = True
    
    criteria3 = [
        '✅ Использована точная формула из документа',
        '✅ Данные правильно подставлены',
        '✅ ВЫПОЛНЕНО арифметическое вычисление',
        '✅ Результат интерпретирован (сравнение с нормами)',
        '❌ НЕТ отказа от вычисления ("я не могу считать")'
    ]
    
    for criterion in criteria3:
        doc.add_paragraph(criterion, style='List Bullet')
    
    doc.add_paragraph()
    
    # Тест 4
    add_heading(doc, 'Тест №4: Неполные данные (проверка уточнения)', 2)
    
    p = doc.add_paragraph()
    p.add_run('Вопрос:').bold = True
    add_code_block(doc, 'Посчитай валовую выручку клиники, если работают 3 врача по 160 часов.')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Ожидаемый ответ (структура):').bold = True
    
    answer_structure4 = """Для расчёта валовой выручки по формуле ВВклиники = Кзаг * НЧ * tраб * n нужны:

Есть данные:
- n = 3 врача
- tраб = 160 часов каждый

Не хватает данных:
- Кзаг (коэффициент загрузки) - доля времени, когда врач занят с пациентами
- НЧ (нормочас) - средний нормочас врачей

Укажите:
1. Коэффициент загрузки в % (например, 80%)
2. Средний нормочас врачей в рублях

По умолчанию для терапевтов:
- Кзаг = 80% (0.8)
- НЧ = 7 500 руб/час

Хотите использовать эти значения?"""
    
    add_code_block(doc, answer_structure4)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Критерии проверки:').bold = True
    
    criteria4 = [
        '✅ Чётко указано, каких данных не хватает',
        '✅ Объяснено, зачем нужна каждая переменная',
        '✅ Предложены значения по умолчанию из норм',
        '❌ НЕТ расчёта с выдуманными данными',
        '❌ НЕТ отказа без объяснения'
    ]
    
    for criterion in criteria4:
        doc.add_paragraph(criterion, style='List Bullet')
    
    doc.add_page_break()
    
    # Тест 5
    add_heading(doc, 'Тест №5: Практическая задача', 2)
    
    p = doc.add_paragraph()
    p.add_run('Вопрос:').bold = True
    add_code_block(doc, 'Что делать, если низкая конверсия из консультации в лечение?')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Ожидаемый ответ (структура):').bold = True
    
    answer_structure5 = """[Определение показателя из документа]

Формула для диагностики:
ОТдок = начатые лечения / проведенные консультации

Шаги решения:
1. Измерить текущий ОТдок по формуле
2. Сравнить с нормой: [норма из документа]
3. Если ниже нормы, проверить:
   - [конкретный фактор 1]
   - [конкретный фактор 2]
   - [конкретный фактор 3]
4. Внедрить контроль: [конкретная рекомендация]

Вопросы:
1. Какие факторы влияют на ОТдок?
2. Как обучить врачей увеличивать конверсию?"""
    
    add_code_block(doc, answer_structure5)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Критерии проверки:').bold = True
    
    criteria5 = [
        '✅ Дано определение проблемы',
        '✅ Указана формула для диагностики',
        '✅ Предложены КОНКРЕТНЫЕ шаги решения (не общие слова)',
        '✅ Упомянуты нормы/целевые значения',
        '❌ НЕТ абстрактных советов без формул',
        '❌ НЕТ рекомендаций, не основанных на документе'
    ]
    
    for criterion in criteria5:
        doc.add_paragraph(criterion, style='List Bullet')
    
    doc.add_page_break()
    
    # 6. Метрики
    add_heading(doc, '6. Метрики для оценки качества', 1)
    
    p = doc.add_paragraph()
    p.add_run('Для методолога:').bold = True
    doc.add_paragraph()
    
    metrics = [
        ('Точность (Accuracy):', 'ответ содержит только информацию из документов (0-100%)'),
        ('Полнота (Completeness):', 'ответ охватывает все аспекты вопроса (0-100%)'),
        ('Формат (Format):', 'ответ следует структуре без markdown (ДА/НЕТ)'),
        ('Вычисления (Calculations):', 'при наличии данных выполнены расчёты (ДА/НЕТ)'),
        ('Уточнения (Follow-up):', 'наличие 2-3 релевантных вопросов (ДА/НЕТ)'),
        ('Галлюцинации (Hallucinations):', 'количество выдуманных фактов (0 = отлично)')
    ]
    
    for i, (key, value) in enumerate(metrics, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {key}').bold = True
        p.add_run(f' {value}')
    
    doc.add_paragraph()
    
    # 7. Шаблон
    add_heading(doc, '7. Шаблон для тестирования', 1)
    
    template = """ТЕСТ №___: [Тип вопроса]
Категория: [Определение / Формула / Расчёт / Практика]
Документ: [Справочник / Золотой Стандарт / Директор]

ВОПРОС:
[Текст вопроса]

ОЖИДАЕМЫЕ ЭЛЕМЕНТЫ ОТВЕТА:
□ Определение/Формула из документа
□ Расшифровка переменных
□ Нормативные значения
□ Пошаговые вычисления (если применимо)
□ Конкретные рекомендации (если применимо)
□ 2-3 уточняющих вопроса

НЕДОПУСТИМО:
□ Markdown форматирование (**, ##, *)
□ Фразы "в документе", "согласно контексту"
□ Выдуманные данные или нормы
□ Отказ от вычислений при наличии данных

ОЦЕНКА:
Точность: ___/100
Полнота: ___/100
Формат: ДА/НЕТ
Вычисления: ДА/НЕТ/N/A
Уточнения: ДА/НЕТ
Галлюцинации: ___

ПРИМЕЧАНИЯ:
[Комментарии методолога]"""
    
    add_code_block(doc, template)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Заключение
    p = doc.add_paragraph()
    p.add_run('Этот стандарт позволит систематически проверять качество работы RAG-системы на разных типах вопросов.').italic = True
    
    # Сохранение
    output_path = 'E:\\СТОМПРАКТИКА ПРОЕКТЫ\\Docling\\methodology\\Стандарт_Вопрос-Ответ_RAG_Тестирование.docx'
    doc.save(output_path)
    print(f'✅ Документ создан: {output_path}')
    
    return output_path

if __name__ == '__main__':
    try:
        create_document()
    except Exception as e:
        print(f'❌ Ошибка: {e}')
        import traceback
        traceback.print_exc()
